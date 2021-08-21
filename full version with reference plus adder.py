import docx
from docx import Document

#import docx2txt

#my_text = docx2txt.process("test.docx")
#document = Document()

doc = Document("manuscript Submission 31 july.docx")
#doc = Document('https://docs.google.com/document/d/1_qhTp5vaNQeWdt7bBNhWThhdA6r3BB1zN41_4Mhzco8/edit')

#print (type(my_text))
#line = list(p.text.split(" "))
print("enter starting value")
a=int(input())
print("increment by ")
add= int(input())
#new = ""
for p in doc.paragraphs:
        line = list(p.text.split(" "))
        for t in range(len(line)):
                        x=line[t]

                        if '[' in x and any(map(str.isalpha, x))==False:
                                    
                                    #z = x.replace("\n", " ")
                                    if '.' in x:
                                        z = x[1:-2]
                                        
                                        if ',' in z:
                                            if ']' in z:
                                                z = z.replace("[", "")
                                                z = z.replace("]", "")
                                            li= z.split(',')
                                            for i in range(len(li)):
                                                b=int(li[i])
                                                if b>a:
                                                        b=b+add
                                                        
                                                li[i]=b
                                                   
                                            li=str(li).replace(" ","")
                                            new =  li+'.' + " "
                                            line[t]= new
                                            p.text=' '.join(map(str, line))
                                            #p.text = p.text.replace(x,new)


                                        elif "–" in z or "-" in z:
                                                  if ']' in z:
                                                            z = z.replace("[", "")
                                                            z = z.replace("]", "")
                                                  z = z.replace("–", "-")
                                                                            
                                                  li= z.split("-")
                                                  for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                      b=b+add
                                                                      
                                                                    
                                                            li[i]=b
                                                  li=str(li).replace(" ","")
                                                  li=str(li).replace(",","-")
                                                  new =  li + "." + " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)
                                              
                                                                  
                                        else:
                                            b=int(z)
                                            if b>a:
                                                    
                                                b=b+add
                                                
                                            new = '['+str(b)+']'+ "." + " "
                                            
                                            line[t]= new
                                            #p.text = p.text.replace(x,new)
                                            p.text=' '.join(map(str, line))


                                    elif '.' not in x and x[-1] != ',' and x[-1] != ';'and x[-1] != ':':
                                              z = x[1:-1]
                        

                                              if ',' in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                    
                                                                b=b+add
                                                                
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li + " "
                                                        #p.text = p.text.replace(x,new)
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        

                                                        
                                              elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")
                                                                            
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+add
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li + " "
                                                        
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)
                                              
                                                                  
                                              else:
                                                  b=int(z)
                                                  if b>a:
                                                          
                                                      b=b+add
                                                      
                                                  new = '['+str(b)+']'+ " "
                                                  
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)

                                                                  

                                    elif '.' not in x and x[-1] == ',':
                                              z = x[1:-2]

                                              if ',' in z:
                                                        if ']' in x:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                    
                                                                b=b+add
                                                                
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li+',' + " "

                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)

                                            

                                              elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")                                              
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+add
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li+',' + " "

                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)


                                              else:
                                                  b=int(z)
                                                  if b>a:
                                                          
                                                      b=b+add
                                                      
                                                  new = '['+str(b)+']'+','+ " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)



                                    elif '.' not in x and x[-1] == ':':
                                              z = x[1:-2]

                                              if ',' in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:

                                                                    
                                                                b=b+add
                                                                
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li+':' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)

                                            

                                              elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")                                              
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+add
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li+':' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)


                                              else:
                                                  b=int(z)
                                                  if b>a:
                                                          
                                                      b=b+add
                                                      
                                                  new = '['+str(b)+']'+':'+ " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)
                                    
                                    elif '.' not in x and x[-1] == ';':
                                          z = x[1:-2]
                                          if ',' in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                    
                                                                b=b+add
                                                                
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li+';' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)

                                            
                           
                                          elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+add
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li+';' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)


                                          else:
                                                  b=int(z)
                                                  if b>a:
                                                          
                                                      b=b+add
                                                      
                                                  new = '['+str(b)+']'+';'+ " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)
                                            
                                                                  
        #p.text=' '.join(map(str, line))                                                


                                        

                        #else:
                            #new += x+" "


#document.add_paragraph(new)
#document.save('test-output.docx')
doc.save('result1.docx')
#doc.save('https://docs.google.com/document/d/1smTUfy1H1NRKhgS0wSAa5MLpsfbimbR4pv6pk3PFYRQ/edit')


            

