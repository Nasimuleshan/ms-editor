import docx
from docx import Document

#import docx2txt

#my_text = docx2txt.process("test.docx")
#document = Document()

doc = Document("test.docx")


#print (type(my_text))
#line = list(p.text.split(" "))
print("enter value")
a=int(input())
#new = ""
for p in doc.paragraphs:
        line = list(p.text.split(" "))
        for t in range(len(line)):
                        x=line[t]

                        if '[' in x and any(map(str.isdigit, x))==True:
                                    
                                    #z = x.replace("\n", " ")
                                    if '.' in x:
                                        z = x[1:-2]
                                        
                                        if ',' in z:
                                            if ']' in z:
                                                z = z.replace("[", "")
                                                z = z.replace("]", "")
                                            li= z.split(',')
                                            for i in range(len(li)):
                                                b=int(li[i])
                                                if b>a:
                                                        b=b+1
                                                li[i]=b
                                                   
                                            li=str(li).replace(" ","")
                                            new =  li+'.' + " "
                                            line[t]= new
                                            p.text=' '.join(map(str, line))
                                            #p.text = p.text.replace(x,new)


                                        elif "–" in z or "-" in z:
                                                  if ']' in z:
                                                            z = z.replace("[", "")
                                                            z = z.replace("]", "")
                                                  z = z.replace("–", "-")
                                                                            
                                                  li= z.split("-")
                                                  for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                      b=b+1
                                                                    
                                                            li[i]=b
                                                  li=str(li).replace(" ","")
                                                  li=str(li).replace(",","-")
                                                  new =  li + "." + " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)
                                              
                                                                  
                                        else:
                                            b=int(z)
                                            if b>a:
                                                b=b+1
                                            new = '['+str(b)+']'+ "." + " "
                                            
                                            line[t]= new
                                            #p.text = p.text.replace(x,new)
                                            p.text=' '.join(map(str, line))


                                    elif '.' not in x and x[-1] != ',' and x[-1] != ';'and x[-1] != ':':
                                              z = x[1:-1]
                        

                                              if ',' in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                b=b+1
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li + " "
                                                        #p.text = p.text.replace(x,new)
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        

                                                        
                                              elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")
                                                                            
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+1
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li + " "
                                                        
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)
                                              
                                                                  
                                              else:
                                                  b=int(z)
                                                  if b>a:
                                                      b=b+1
                                                  new = '['+str(b)+']'+ " "
                                                  
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)

                                                                  

                                    elif '.' not in x and x[-1] == ',':
                                              z = x[1:-2]

                                              if ',' in z:
                                                        if ']' in x:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                b=b+1
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li+',' + " "

                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)

                                            

                                              elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")                                              
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+1
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li+',' + " "

                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)


                                              else:
                                                  b=int(z)
                                                  if b>a:
                                                      b=b+1
                                                  new = '['+str(b)+']'+','+ " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)



                                    elif '.' not in x and x[-1] == ':':
                                              z = x[1:-2]

                                              if ',' in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                b=b+1
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li+':' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)

                                            

                                              elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")                                              
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+1
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li+':' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)


                                              else:
                                                  b=int(z)
                                                  if b>a:
                                                      b=b+1
                                                  new = '['+str(b)+']'+':'+ " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)
                                    
                                    elif '.' not in x and x[-1] == ';':
                                          z = x[1:-2]
                                          if ',' in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        li= z.split(',')
                                                        for i in range(len(li)):
                                                            b=int(li[i])
                                                            if b>a:
                                                                b=b+1
                                                            li[i]=b
                                                        li=str(li).replace(" ","")
                                                        new =  li+';' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)

                                            
                           
                                          elif "–" in z or "-" in z:
                                                        if ']' in z:
                                                                  z = z.replace("[", "")
                                                                  z = z.replace("]", "")
                                                        z = z.replace("–", "-")
                                                        li= z.split("-")
                                                        for i in range(len(li)):
                                                                  b=int(li[i])
                                                                  if b>a:
                                                                            b=b+1
                                                                            
                                                                  li[i]=b
                                                        li=str(li).replace(" ","")
                                                        li=str(li).replace(",","-")
                                                        new =  li+';' + " "
                                                        line[t]= new
                                                        p.text=' '.join(map(str, line))
                                                        #p.text = p.text.replace(x,new)


                                          else:
                                                  b=int(z)
                                                  if b>a:
                                                      b=b+1
                                                  new = '['+str(b)+']'+';'+ " "
                                                  line[t]= new
                                                  p.text=' '.join(map(str, line))
                                                  #p.text = p.text.replace(x,new)
                                            
                                                                  
        #p.text=' '.join(map(str, line))                                                


                                        

                        #else:
                            #new += x+" "


#document.add_paragraph(new)
#document.save('test-output.docx')
doc.save('result1.docx')


            

